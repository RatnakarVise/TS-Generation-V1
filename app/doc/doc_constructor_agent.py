import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client as win32
import tempfile
import shutil
import time
import multiprocessing

def safe_update_index(in_path, out_path):
    return update_index_with_word(in_path, out_path)

def add_heading(doc, text, level):
    para = doc.add_heading(text, level=level)
    return para if para is not None else doc.add_paragraph(text, style="Heading1")
    # doc.add_heading(text, level=level)

def add_table(doc, colnames, rows):
    table = doc.add_table(rows=1, cols=len(colnames))
    table.style = "Light List"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(colnames):
        hdr_cells[i].text = str(col)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val) if val is not None else ""
    return table

def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)

def add_bookmark(paragraph, name, bid):
    """Create a collapsed bookmark on the given paragraph."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)


    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))


    # Insert start at beginning of paragraph, end at end
    paragraph._p.insert(0, start)
    paragraph._p.append(end)

def add_pageref_field(paragraph, bookmark_name):
    """Insert a PAGEREF field that will display the page of the bookmark."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); 
    fldBegin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText"); 
    instr.set(qn("xml:space"), "preserve"); 
    instr.text = f"PAGEREF {bookmark_name} \\h"

    fldSep = OxmlElement("w:fldChar"); 
    fldSep.set(qn("w:fldCharType"), "separate")

    result = OxmlElement("w:t")
    result.text = "1"   # dummy, will be replaced by Word after update

    fldEnd = OxmlElement("w:fldChar"); 
    fldEnd.set(qn("w:fldCharType"), "end")

    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(result)
    run._r.append(fldEnd)
    
def find_section_content(content_list, section_title):
    for sec in content_list:
        if sec.get('section_name', '').lower().strip() == section_title.lower().strip():
            return sec['content']
    return None

def find_all_markdown_tables_and_text(text):
    if not text:
        return []
    table_pattern = re.compile(
        r'((?:\|.*\n)+?\|[ \t\-\|:]+\|\n(?:\|.*\n?)+)',
        re.MULTILINE
    )
    chunks = []
    last_idx = 0
    for match in table_pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            txt = text[last_idx:start].strip()
            if txt:
                chunks.append(('text', txt))
        table_md = match.group(1).strip()
        chunks.append(('table', table_md))
        last_idx = end
    if last_idx < len(text):
        txt = text[last_idx:].strip()
        if txt:
            chunks.append(('text', txt))
    return chunks

def parse_markdown_table(table_md):
    lines = [line.strip() for line in table_md.strip().splitlines() if line.strip() and line.strip().startswith('|')]
    if not lines:
        return None, None
    rows = [[cell.strip() for cell in l.strip('|').split('|')] for l in lines]
    if len(rows) < 2:
        return None, None
    divider_row = rows[1]
    if all(re.match(r'^[-:\s]+$', c) for c in divider_row):
        del rows[1]
    colnames = rows[0]
    data_rows = rows[1:]
    return colnames, data_rows

# This is your robust extractor for the flow diagram agent:
def extract_arrow_flow(text):
    if not text:
        return ""
    for line in text.splitlines():
        line = line.strip("` ").strip()
        if "->" in line and not line.lower().startswith(('diagram', 'flow', 'legend', '#')):
            return line
    if "->" in text:
        return text.strip()
    return ""

def build_document(content, sections,out_path="TechSpec.docx", flow_diagram_agent=None, diagram_dir="diagrams"):
    doc = Document()
    
    
    # Add the main heading at the top
    add_heading(doc, "Technical Specification Document", 0)

    # Precompute bookmark names (used by index + later applied to headings)
    bookmark_names = [f"sec_{i+1}" for i in range(len(sections))]
    add_heading(doc, "Index", 1)

    # 2) Custom Index on page 1 with PAGEREF fields (real page numbers after Update Field)
    # idx_heading = add_heading(doc, "Index", 1)

    # Set a right-aligned tab with dot leader for each index entry (looks like: Title ....... 3)
    right_pos = Inches(6.0)  # adjust if your page margins differ
    for i, section in enumerate(sections):
        title = section.get("title")
        p = doc.add_paragraph()
        # Tab stops for this paragraph
        p.paragraph_format.tab_stops.add_tab_stop(
            right_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        p.add_run(f"{i+1}. {title}")
        p.add_run("\t")  # jump to the right-aligned tab stop
        add_pageref_field(p, bookmark_names[i])  # page number field pointing to the heading's bookmark

    doc.add_page_break()

    p = doc.add_paragraph()
    add_toc(p)
    # doc.add_page_break() 
    for i, section in enumerate(sections):
        title = section.get("title")
        header = f"{i+1}. {title}"
        h = add_heading(doc, header, 1)
        add_bookmark(h, bookmark_names[i], i+1)  # <-- anchor for PAGEREF

        sec_content = find_section_content(content, title)

        # FLOW DIAGRAM SECTION HANDLING
        if title.strip().lower() == "flow diagram":
            diagram_img = None
            if flow_diagram_agent is not None and sec_content:
                try:
                    flow_line = extract_arrow_flow(sec_content)
                    if flow_line:
                        diagram_img = flow_diagram_agent.run(flow_line)  # <-- Returns BytesIO
                    else:
                         diagram_img = None
                except Exception as e:
                    print(f"Flow diagram agent error: {e}")
                    diagram_img = None
            if diagram_img:
                 doc.add_picture(diagram_img, width=Inches(5.5))
            else:
                 doc.add_paragraph("[Flow diagram not available]")
            continue   # Skip remaining processing for this section

        # Universal parsing for text+tables:
        chunks = find_all_markdown_tables_and_text(sec_content)
        for typ, value in chunks:
            if typ == 'text':
                if value:
                    doc.add_paragraph(value)
            elif typ == 'table':
                colnames, rows = parse_markdown_table(value)
                if colnames and rows:
                    add_table(doc, colnames, rows)
                else:
                    doc.add_paragraph(value)

    doc.add_paragraph("\nDocument generated by PWC AI-powered ABAP Tech Spec Assistant.")
    # Save to temporary file
    # tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    # os.close(tmp_fd)
    # doc.save(tmp_path)

    # # Update fields and page numbers in Word
    # # final_path = update_index_with_word(tmp_path, out_path)
    # with multiprocessing.Pool(1) as pool:
    #     final_path = pool.apply(safe_update_index, args=(tmp_path, out_path))

    # if os.path.exists(tmp_path):
    #     os.remove(tmp_path)

    # return final_path
    return doc
def update_index_with_word(in_path, out_path=None, max_retries=5, retry_delay=1):
    """
    Open a Word document via COM, update all fields (TOC, PAGEREF, headers/footers), 
    save to out_path, and close Word safely.

    Args:
        in_path (str): Input Word file path.
        out_path (str): Output Word file path. If None, overwrites input.
        max_retries (int): Number of retries if Word is busy.
        retry_delay (float): Seconds to wait between retries.

    Returns:
        str: The absolute path to the saved Word document.
    """
    word = None
    try:
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone

        # Retry opening document if Word is busy
        for attempt in range(max_retries):
            try:
                doc = word.Documents.Open(os.path.abspath(in_path))
                break
            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                else:
                    raise RuntimeError(f"Failed to open Word file after {max_retries} attempts: {e}")

        # Repaginate and update all fields
        doc.Repaginate()
        doc.Fields.Update()

        # Update all Table of Contents
        for toc in doc.TablesOfContents:
            toc.Update()

        # Update fields in headers and footers
        for section in doc.Sections:
            for header in section.Headers:
                header.Range.Fields.Update()
            for footer in section.Footers:
                footer.Range.Fields.Update()

        # Update footnotes and endnotes
        try: doc.Footnotes.Range.Fields.Update()
        except: pass
        try: doc.Endnotes.Range.Fields.Update()
        except: pass

        # Save document
        final_path = os.path.abspath(out_path) if out_path else os.path.abspath(in_path)
        doc.SaveAs(final_path)
        doc.Close(SaveChanges=True)

        return final_path

    finally:
        if word:
            try:
                word.Quit()
            except:
                pass
            del word